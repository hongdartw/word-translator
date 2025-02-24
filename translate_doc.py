import os
from docx import Document
from openai import OpenAI
from dotenv import load_dotenv
from pathlib import Path
from ai_settings import AIProvider, AISettings
import google.generativeai as genai
import keyboard  # 新增 keyboard 套件
import sys
from threading import Event
import time

# 修改環境變數和路徑設定
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ENV_PATH = os.path.join(BASE_DIR, '.env')
INPUT_DIR = os.path.join(BASE_DIR, 'input')
OUTPUT_DIR = os.path.join(BASE_DIR, 'output')

# 載入環境變數
load_dotenv(ENV_PATH)

# 建立停止事件
stop_event = Event()

def check_esc_pressed():
    """
    檢查是否按下 ESC 鍵
    """
    if keyboard.is_pressed('esc'):
        print("\n\n程序已被使用者中斷！")
        stop_event.set()
        return True
    return False

def get_ai_client(provider: AIProvider):
    """
    根據選擇的提供者返回對應的 AI 客戶端
    """
    settings = AISettings.get_api_settings(provider)

    if provider == AIProvider.GEMINI:
        genai.configure(api_key=settings["api_key"])
        return genai.GenerativeModel(settings["model"])
    else:
        return OpenAI(
            api_key=settings["api_key"],
            base_url=settings["api_url"]
        )

def translate_text(text, target_language, provider: AIProvider, ai_client):
    """
    使用選定的 AI 服務翻譯文本
    """
    if text.startswith('http://') or text.startswith('https://'):
        return text

    try:
        prompt = f"""
        Translate the following text to {target_language}.
        Important rules:
        1. Keep it simple and direct
        2. Do not add any explanatory text
        3. Just provide the translation

        Text to translate: {text}
        """

        if provider == AIProvider.GEMINI:
            response = ai_client.generate_content(prompt)
            return response.text.strip()
        else:
            response = ai_client.chat.completions.create(
                model=AISettings.get_api_settings(provider)["model"],
                messages=[
                    {"role": "system", "content": "You are a translator. Provide direct translations without additional explanations."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=1000
            )
            return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"翻譯時發生錯誤: {str(e)}")
        return text

def translate_filename(filename, target_language, provider: AIProvider, ai_client):
    """
    翻譯檔案名稱（不包含副檔名）
    """
    name_without_ext = os.path.splitext(filename)[0]
    translated_name = translate_text(name_without_ext, target_language, provider, ai_client)
    return f"{translated_name}.docx"

def process_document(input_file, target_language, provider: AIProvider, ai_client):
    """
    處理Word文件並進行翻譯，保留所有格式和圖片
    """
    try:
        doc = Document(input_file)
        output_filename = translate_filename(os.path.basename(input_file), target_language, provider, ai_client)
        output_path = os.path.join(OUTPUT_DIR, output_filename)

        # 定義 Word 文檔的 XML 命名空間
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main'
        }

        def has_picture(run):
            """檢查 run 是否包含圖片"""
            try:
                drawing = run._element.xpath('.//w:drawing | .//w:pict', namespaces=namespaces)
                return bool(drawing)
            except Exception:
                return False

        # 處理文檔內容
        for section in doc.sections:
            if stop_event.is_set():
                return
            if check_esc_pressed():
                return

            # 處理頁首
            header = section.header
            if not header.is_linked_to_previous:
                for paragraph in header.paragraphs:
                    if paragraph.text.strip():
                        # 保存原始的 run 格式和內容
                        original_runs = []
                        for run in paragraph.runs:
                            original_runs.append({
                                'text': run.text,
                                'bold': run.bold,
                                'italic': run.italic,
                                'underline': run.underline,
                                'font': run.font.name,
                                'size': run.font.size,
                                'color': run.font.color.rgb if run.font.color else None,
                                'highlight_color': run.font.highlight_color,
                                'has_picture': has_picture(run),
                                '_element': run._element
                            })

                        # 收集需要翻譯的文字
                        text_to_translate = ""
                        text_runs = []
                        for run_format in original_runs:
                            if not run_format['has_picture']:
                                text_to_translate += run_format['text']
                                text_runs.append(run_format)

                        # 只有在有文字需要翻譯時才進行翻譯
                        if text_to_translate.strip():
                            translated_text = translate_text(text_to_translate.strip(), target_language, provider, ai_client)

                            # 清空段落但保留格式
                            paragraph.clear()

                            # 重建段落，保留原始格式和圖片
                            current_pos = 0
                            total_text_length = len(text_to_translate.strip())

                            for run_format in original_runs:
                                new_run = paragraph.add_run()

                                if run_format['has_picture']:
                                    # 複製原始 run 的所有內容，包括圖片
                                    new_run._element.append(run_format['_element'])
                                else:
                                    # 處理文字
                                    if total_text_length > 0:
                                        text_length = int(len(translated_text) * (len(run_format['text'].strip()) / total_text_length))
                                        if text_length > 0:
                                            new_run.text = translated_text[current_pos:current_pos + text_length]
                                            current_pos += text_length

                                # 恢復格式
                                new_run.bold = run_format['bold']
                                new_run.italic = run_format['italic']
                                new_run.underline = run_format['underline']
                                new_run.font.name = run_format['font']
                                new_run.font.size = run_format['size']
                                if run_format['color']:
                                    new_run.font.color.rgb = run_format['color']
                                new_run.font.highlight_color = run_format['highlight_color']

            # 處理頁尾
            footer = section.footer
            if not footer.is_linked_to_previous:
                # 翻譯頁尾段落
                for paragraph in footer.paragraphs:
                    if paragraph.text.strip():
                        translated_text = translate_text(paragraph.text, target_language, provider, ai_client)
                        paragraph.text = translated_text

                # 翻譯頁尾表格
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                translated_text = translate_text(cell.text, target_language, provider, ai_client)
                                cell.text = translated_text

        # 翻譯段落
        for paragraph in doc.paragraphs:
            if stop_event.is_set():
                return
            if check_esc_pressed():
                return

            if paragraph.text.strip():
                # 保存原始的 run 格式和內容
                original_runs = []
                for run in paragraph.runs:
                    original_runs.append({
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font': run.font.name,
                        'size': run.font.size,
                        'color': run.font.color.rgb if run.font.color else None,
                        'highlight_color': run.font.highlight_color,
                        'has_picture': has_picture(run),
                        '_element': run._element
                    })

                # 收集需要翻譯的文字
                text_to_translate = ""
                text_runs = []
                for run_format in original_runs:
                    if not run_format['has_picture']:
                        text_to_translate += run_format['text']
                        text_runs.append(run_format)

                # 只有在有文字需要翻譯時才進行翻譯
                if text_to_translate.strip():
                    translated_text = translate_text(text_to_translate.strip(), target_language, provider, ai_client)

                    # 清空段落但保留格式
                    paragraph.clear()

                    # 重建段落，保留原始格式和圖片
                    current_pos = 0
                    total_text_length = len(text_to_translate.strip())

                    for run_format in original_runs:
                        new_run = paragraph.add_run()

                        if run_format['has_picture']:
                            # 複製原始 run 的所有內容，包括圖片
                            new_run._element.append(run_format['_element'])
                        else:
                            # 處理文字
                            if total_text_length > 0:
                                text_length = int(len(translated_text) * (len(run_format['text'].strip()) / total_text_length))
                                if text_length > 0:
                                    new_run.text = translated_text[current_pos:current_pos + text_length]
                                    current_pos += text_length

                        # 恢復格式
                        new_run.bold = run_format['bold']
                        new_run.italic = run_format['italic']
                        new_run.underline = run_format['underline']
                        new_run.font.name = run_format['font']
                        new_run.font.size = run_format['size']
                        if run_format['color']:
                            new_run.font.color.rgb = run_format['color']
                        new_run.font.highlight_color = run_format['highlight_color']

        # 翻譯表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if stop_event.is_set():
                        return
                    if check_esc_pressed():
                        return
                    if cell.text.strip():
                        translated_text = translate_text(cell.text, target_language, provider, ai_client)
                        cell.text = translated_text

        if stop_event.is_set():
            return

        # 保存文件
        doc.save(output_path)
        print(f"文件已保存為: {output_filename}")

    except Exception as e:
        print(f"處理文件時發生錯誤: {str(e)}")
        import traceback
        traceback.print_exc()  # 印出詳細的錯誤訊息

def test_api_response_time(provider: AIProvider):
    """
    測試 API 的回應時間
    """
    try:
        ai_client = get_ai_client(provider)
        test_text = "Hello, this is a test message."

        start_time = time.time()
        translate_text(test_text, "chinese", provider, ai_client)
        end_time = time.time()

        response_time = round((end_time - start_time) * 1000)  # 轉換為毫秒
        return True, response_time
    except Exception as e:
        return False, None

def test_all_apis():
    """
    測試所有可用的 API 接口及其回應時間
    """
    available_providers = {
        "Grok": AIProvider.GROK,
        "Free ChatGPT": AIProvider.FREE_CHATGPT,
        "Gemini": AIProvider.GEMINI
    }

    working_providers = []
    print("\n正在測試所有 API 接口...")

    for name, provider in available_providers.items():
        print(f"\n測試 {name} API...")
        success, message = AISettings.test_api_connection(provider)
        print(message)

        if success:
            print("測試回應時間中...")
            time_success, response_time = test_api_response_time(provider)
            if time_success:
                print(f"回應時間: {response_time} 毫秒")
                working_providers.append((name, provider, response_time))
            else:
                print("回應時間測試失敗")
                working_providers.append((name, provider, float('inf')))

    # 根據回應時間排序
    working_providers.sort(key=lambda x: x[2])
    return working_providers

def main():
    """
    主程序
    """
    print("\n歡迎使用文件翻譯工具")
    print("提示：按 ESC 鍵可隨時中斷程序")

    try:
        # 確保輸出目錄存在
        Path(OUTPUT_DIR).mkdir(parents=True, exist_ok=True)

        print("\n首先進行 API 可用性測試")

        # 測試所有 API
        working_providers = test_all_apis()

        if not working_providers:
            print("\n錯誤：沒有可用的 API 接口")
            return

        # 顯示可用的 API 選項（包含回應時間）
        print("\n可用的 API 服務（按回應時間排序）：")
        for i, (name, _, response_time) in enumerate(working_providers, 1):
            if response_time == float('inf'):
                print(f"{i}. {name} (回應時間測試失敗)")
            else:
                print(f"{i}. {name} (回應時間: {response_time}ms)")

        # 選擇 API 提供者
        while True:
            if check_esc_pressed():
                return
            provider_choice = input("\n請選擇要使用的 API (輸入數字): ").strip()
            try:
                choice_idx = int(provider_choice) - 1
                if 0 <= choice_idx < len(working_providers):
                    selected_provider = working_providers[choice_idx][1]
                    print(f"\n已選擇: {working_providers[choice_idx][0]}")
                    break
                else:
                    print("無效的選擇，請重新輸入")
            except ValueError:
                print("請輸入有效的數字")

        if stop_event.is_set():
            return

        # 初始化 AI 客戶端
        ai_client = get_ai_client(selected_provider)

        # 獲取輸入目錄中的所有.docx文件
        docx_files = [f for f in os.listdir(INPUT_DIR) if f.endswith('.docx')]

        if not docx_files:
            print("在輸入目錄中未找到.docx文件")
            return

        # 顯示可用的文件
        print("\n可用的文件：")
        for i, file in enumerate(docx_files, 1):
            print(f"{i}. {file}")

        # 選擇目標語言
        print("\n選擇目標語言：")
        print("1. 英文 (English)")
        print("2. 泰文 (Thai)")

        if stop_event.is_set():
            return

        language_choice = input("請輸入選項 (1 或 2): ").strip()
        target_language = "english" if language_choice == "1" else "thai"

        # 處理每個文件
        for file in docx_files:
            if stop_event.is_set():
                break
            input_path = os.path.join(INPUT_DIR, file)
            print(f"\n正在處理文件: {file}")
            process_document(input_path, target_language, selected_provider, ai_client)

    except KeyboardInterrupt:
        print("\n\n程序已被使用者中斷！")
    finally:
        if stop_event.is_set():
            print("\n程序已終止")
            sys.exit(0)

if __name__ == "__main__":
    main()